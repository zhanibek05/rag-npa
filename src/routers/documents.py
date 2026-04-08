"""
Роутер для управления документами:
  GET    /documents          — список с пагинацией
  POST   /documents/upload   — загрузить DOCX / PDF / TXT
  DELETE /documents/{id}     — удалить из Postgres + Qdrant
"""

import hashlib
import io
import re
import uuid
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from qdrant_client import QdrantClient
from qdrant_client.models import Filter, FieldCondition, MatchValue, FilterSelector
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

try:
    from src.core.config import QDRANT_URL, QDRANT_COLLECTION, EMBEDDING_MODEL, EMBEDDING_DEVICE
    from src.core.database import get_db
    from src.core.models import Document
except ModuleNotFoundError:
    from core.config import QDRANT_URL, QDRANT_COLLECTION, EMBEDDING_MODEL, EMBEDDING_DEVICE
    from core.database import get_db
    from core.models import Document

router = APIRouter(prefix="/documents", tags=["documents"])


# ── Pydantic schemas ──────────────────────────────────────────────────────────

class DocumentOut(BaseModel):
    id: str
    title: Optional[str]
    doc_type: Optional[str]
    status: Optional[str]
    adopted_date: Optional[str]
    source: str
    index_status: str
    url: Optional[str]

    class Config:
        from_attributes = True


class DocumentsPage(BaseModel):
    items: List[DocumentOut]
    total: int
    page: int
    page_size: int


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n".join(pages)


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    return _extract_txt(data)


# ── Chunking + indexing ───────────────────────────────────────────────────────

def _chunk_text(text: str, max_chars: int = 1200, overlap: int = 2) -> list[dict]:
    paragraphs = [p for p in re.split(r"\n+", text) if p.strip()]
    chunks, i = [], 0
    while i < len(paragraphs):
        buf, length, start = [], 0, i
        while i < len(paragraphs):
            p = paragraphs[i]
            if buf and length + len(p) + 1 > max_chars:
                break
            buf.append(p)
            length += len(p) + 1
            i += 1
        chunks.append({
            "text": "\n".join(buf),
            "chunk_index": len(chunks),
            "paragraph_start": start,
            "paragraph_end": i - 1,
        })
        if overlap > 0:
            i = max(i - overlap, start + 1)
    return chunks


def _point_id(doc_id: str, chunk_index: int) -> int:
    raw = f"{doc_id}_{chunk_index}".encode()
    return int(hashlib.md5(raw).hexdigest()[:13], 16)


def _index_document(doc: Document) -> int:
    """Чанкует и индексирует документ в Qdrant. Возвращает кол-во чанков."""
    from sentence_transformers import SentenceTransformer
    from qdrant_client.models import PointStruct, VectorParams, Distance

    chunks = _chunk_text(doc.raw_text or "")
    if not chunks:
        return 0

    model = SentenceTransformer(EMBEDDING_MODEL, device=EMBEDDING_DEVICE)
    texts = [c["text"] for c in chunks]
    vectors = model.encode(texts, normalize_embeddings=True, convert_to_numpy=True)

    qdrant = QdrantClient(url=QDRANT_URL)
    if not qdrant.collection_exists(QDRANT_COLLECTION):
        dim = vectors.shape[1]
        qdrant.create_collection(
            collection_name=QDRANT_COLLECTION,
            vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
        )

    points = [
        PointStruct(
            id=_point_id(doc.id, c["chunk_index"]),
            vector=vectors[j].tolist(),
            payload={
                "doc_id": doc.id,
                "text": c["text"],
                "title": doc.title,
                "doc_type": doc.doc_type,
                "status": doc.status,
                "source": doc.source,
                "source_url": doc.url,
                "chunk_index": c["chunk_index"],
            },
        )
        for j, c in enumerate(chunks)
    ]
    qdrant.upsert(collection_name=QDRANT_COLLECTION, points=points)
    return len(points)


def _delete_from_qdrant(doc_id: str) -> None:
    qdrant = QdrantClient(url=QDRANT_URL)
    if not qdrant.collection_exists(QDRANT_COLLECTION):
        return
    qdrant.delete(
        collection_name=QDRANT_COLLECTION,
        points_selector=FilterSelector(
            filter=Filter(
                must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
            )
        ),
    )


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.get("", response_model=DocumentsPage)
async def list_documents(
    page: int = 1,
    page_size: int = 20,
    source: Optional[str] = None,
    doc_type: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
):
    q = select(Document)
    if source:
        q = q.where(Document.source == source)
    if doc_type:
        q = q.where(Document.doc_type == doc_type)

    total = await db.scalar(select(func.count()).select_from(q.subquery()))
    items = (await db.scalars(
        q.order_by(Document.adopted_date.desc().nulls_last())
         .offset((page - 1) * page_size)
         .limit(page_size)
    )).all()

    return DocumentsPage(
        items=[
            DocumentOut(
                id=d.id,
                title=d.title,
                doc_type=d.doc_type,
                status=d.status,
                adopted_date=d.adopted_date.isoformat() if d.adopted_date else None,
                source=d.source,
                index_status=d.index_status,
                url=d.url,
            )
            for d in items
        ],
        total=total,
        page=page,
        page_size=page_size,
    )


@router.post("/upload", response_model=DocumentOut)
async def upload_document(
    title: str = Form(...),
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    allowed = (".docx", ".pdf", ".txt")
    if not any(file.filename.lower().endswith(ext) for ext in allowed):
        raise HTTPException(status_code=400, detail="Поддерживаются только .docx, .pdf, .txt")

    data = await file.read()
    try:
        raw_text = extract_text(file.filename, data)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Не удалось извлечь текст: {e}")

    if len(raw_text.strip()) < 50:
        raise HTTPException(status_code=422, detail="Документ слишком короткий или пустой")

    doc_id = f"custom_{uuid.uuid4().hex[:12]}"
    doc = Document(
        id=doc_id,
        url=f"custom://{doc_id}",
        title=title,
        doc_type="Корпоративный документ",
        status="Действующий",
        raw_text=raw_text,
        scraped_at=datetime.now(timezone.utc),
        index_status="scraped",
        source="custom",
    )
    db.add(doc)
    await db.flush()

    try:
        _index_document(doc)
        doc.index_status = "indexed"
    except Exception as e:
        doc.index_status = "failed"
        await db.commit()
        raise HTTPException(status_code=500, detail=f"Ошибка индексации: {e}")

    await db.commit()
    await db.refresh(doc)

    return DocumentOut(
        id=doc.id,
        title=doc.title,
        doc_type=doc.doc_type,
        status=doc.status,
        adopted_date=None,
        source=doc.source,
        index_status=doc.index_status,
        url=doc.url,
    )


@router.delete("/{doc_id}")
async def delete_document(doc_id: str, db: AsyncSession = Depends(get_db)):
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")

    _delete_from_qdrant(doc_id)
    await db.delete(doc)
    await db.commit()
    return {"ok": True}
